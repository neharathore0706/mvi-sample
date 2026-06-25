from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/neha/Documents/githubprojects/MVI Sample/MVI_Architecture_Medium_Article.docx"

NAVY = RGBColor(25, 55, 78)
BLUE = RGBColor(37, 99, 145)
MUTED = RGBColor(92, 104, 112)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F7"
CODE_BG = "F5F5F5"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(code.strip("\n").splitlines()):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.7, color=RGBColor(35, 35, 35))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=160, start=200, bottom=160, end=200)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=NAVY)
    text_run = p.add_run(text)
    set_run_font(text_run, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.4)
section.footer_distance = Inches(0.4)

# Styles: editorial long-form preset resolved into explicit values.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in (
    ("Heading 1", 17, BLUE, 17, 7),
    ("Heading 2", 13.5, BLUE, 13, 5),
    ("Heading 3", 11.5, NAVY, 9, 4),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

list_style = doc.styles["List Bullet"]
list_style.font.name = "Calibri"
list_style.font.size = Pt(11)
list_style.paragraph_format.left_indent = Inches(0.5)
list_style.paragraph_format.first_line_indent = Inches(-0.25)
list_style.paragraph_format.space_after = Pt(4)
list_style.paragraph_format.line_spacing = 1.15

# Running furniture.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("ANDROID ARCHITECTURE  •  MVI")
set_run_font(hr, size=8.5, color=MUTED, bold=True)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Practical MVI Architecture in Android")
set_run_font(fr, size=8.5, color=MUTED)

# Editorial cover.
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(54)

kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(12)
kr = kicker.add_run("ANDROID ARCHITECTURE")
set_run_font(kr, size=10, color=BLUE, bold=True)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
tr = title.add_run("A Practical Guide to MVI Architecture in Android")
set_run_font(tr, size=28, color=NAVY, bold=True)
set_paragraph_keep(title, keep_next=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(18)
sr = subtitle.add_run(
    "Building predictable Compose screens with intents, immutable state, StateFlow, Retrofit, coroutines, and Hilt"
)
set_run_font(sr, size=14, color=MUTED)

dek = doc.add_paragraph()
dek.alignment = WD_ALIGN_PARAGRAPH.CENTER
dek.paragraph_format.left_indent = Inches(0.45)
dek.paragraph_format.right_indent = Inches(0.45)
dek.paragraph_format.space_after = Pt(22)
dr = dek.add_run(
    "MVI can sound theoretical until you follow one real feature from a tap on the screen to a network request and back to a rendered UI. This article does exactly that with an animal-list application."
)
set_run_font(dr, size=11.5, color=NAVY, italic=True)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Medium-ready technical article  •  Kotlin  •  Jetpack Compose")
set_run_font(mr, size=9.5, color=MUTED)

doc.add_page_break()

add_body(
    doc,
    "Android screens become difficult to maintain when they can be updated from many places. A button changes one variable, a network callback changes another, and navigation introduces a third path. Soon the UI no longer has one clear answer to a simple question: what should I display right now?"
)
add_body(
    doc,
    "Model–View–Intent, usually shortened to MVI, addresses that problem by forcing state changes through a predictable one-way loop. The UI sends an intent. The ViewModel interprets it and performs the required work. The ViewModel then publishes a new immutable state. The UI renders that state."
)

add_callout(
    doc,
    "The core loop",
    "View sends Intent → ViewModel processes Intent → Repository performs data work → ViewModel emits State → View renders State."
)

doc.add_heading("What problems can appear in MVVM?", level=1)
add_body(
    doc,
    "MVVM is not inherently broken, and a well-designed MVVM application can be clean, scalable, and testable. The difficulty is that MVVM defines responsibilities but does not strictly define how every UI action must become a state transition. Without consistent conventions, teams often implement the same pattern in very different ways."
)
add_body(
    doc,
    "The following problems are therefore not guaranteed weaknesses of MVVM. They are common failure modes in loosely structured MVVM implementations:"
)
for text in (
    "Multiple sources of truth: a screen may observe separate values for data, loading, errors, and selection while also keeping local UI state.",
    "Invalid state combinations: independent flags can accidentally describe impossible situations, such as showing loading and an error at the same time.",
    "Unclear action handling: the view may call many unrelated ViewModel methods, making it harder to trace which user action produced the current state.",
    "Scattered state mutation: several methods may update overlapping observable properties, so state transitions are not visible in one place.",
    "One-time event confusion: navigation, snackbars, and toasts are sometimes modeled as persistent state and can be repeated after recreation.",
    "Large ViewModels: networking, validation, mapping, navigation decisions, and state mutation can gradually accumulate in one class.",
    "Inconsistent team implementations: one feature may use LiveData, another StateFlow, another callbacks, and another local Compose state for similar responsibilities.",
):
    add_bullet(doc, text)

add_body(
    doc,
    "For example, a ViewModel might expose animals, isLoading, error, and selectedAnimal as four independent streams. The UI must combine their latest values correctly, and every writer must preserve valid combinations."
)
add_code(
    doc,
    """
val animals: StateFlow<List<Animal>>
val isLoading: StateFlow<Boolean>
val error: StateFlow<String?>
val selectedAnimal: StateFlow<Animal?>
"""
)
add_body(
    doc,
    "MVI responds by applying stricter conventions: user actions enter through intents, the feature exposes one immutable state snapshot, and the UI renders that state. Importantly, modern MVVM can adopt the same practices. In Compose applications, MVVM with an immutable UiState and one-way event handling can look almost identical to practical MVI."
)
add_callout(
    doc,
    "Balanced view",
    "MVI is not a replacement for every MVVM application. It is a stricter way to organize state and actions when predictability matters more than minimizing boilerplate."
)

doc.add_heading("Why MVI works particularly well with Compose", level=1)
add_body(
    doc,
    "Jetpack Compose is declarative: instead of manually commanding individual views to change, we describe what the screen should look like for a given state. MVI supplies exactly that state. A composable collects a StateFlow and chooses the correct UI for loading, content, detail, or error."
)
add_body(
    doc,
    "The result is easier to reason about because the screen is not assembled from scattered flags and callbacks. It is a function of one observable state."
)

doc.add_heading("The sample feature", level=1)
add_body(
    doc,
    "The feature in this article fetches a list of animals from a JSON endpoint. While the request is running, the screen displays rounded shimmer placeholders. When data arrives, it shows a 40 × 40 thumbnail, the animal name, and a three-line description. Selecting a row opens a full-screen detail view with the complete image and description."
)
add_body(
    doc,
    "The implementation uses Retrofit for networking, coroutines for asynchronous work, Hilt for dependency injection, Glide for image loading, and StateFlow for observable UI state."
)

doc.add_heading("1. Model user actions as intents", level=1)
add_body(
    doc,
    "An intent describes what happened. It should express the user’s action or the UI’s request without containing implementation details such as Retrofit calls."
)
add_code(
    doc,
    """
sealed interface MainIntent {
    data object FetchAnimals : MainIntent
    data class SelectAnimal(val animal: Animal) : MainIntent
    data object BackToAnimalList : MainIntent
}
"""
)
add_body(
    doc,
    "These intents form a small public language between the screen and the ViewModel. The UI does not call the repository, decide how to cache data, or mutate state directly."
)

doc.add_heading("2. Keep one source of truth for UI state", level=1)
add_body(
    doc,
    "A common early implementation uses a sealed state for the visible screen and also stores a separate animal list in the ViewModel. That creates two sources of truth. If those values diverge, back navigation can restore stale content."
)
add_body(
    doc,
    "A scalable approach is one immutable state object containing all information needed to render the feature:"
)
add_code(
    doc,
    """
data class MainUiState(
    val animals: List<Animal> = emptyList(),
    val selectedAnimal: Animal? = null,
    val isLoading: Boolean = false,
    val error: String? = null
)
"""
)
add_body(
    doc,
    "The list and detail screens are now two representations of the same state. When selectedAnimal is null, the UI shows the list. When it contains an animal, the UI shows the detail screen. Returning to the list only clears the selection; the fetched data remains in state."
)

doc.add_heading("3. Expose state, but never expose mutation", level=1)
add_body(
    doc,
    "The ViewModel owns MutableStateFlow and exposes only StateFlow. This allows the screen to observe changes while preventing it from assigning arbitrary values."
)
add_code(
    doc,
    """
private val _state = MutableStateFlow(MainUiState())
val state: StateFlow<MainUiState> = _state.asStateFlow()
"""
)
add_callout(
    doc,
    "Rule of thumb",
    "The ViewModel may mutate _state. The UI may only collect state and send intents."
)

doc.add_heading("4. Process intents with a simple function", level=1)
add_body(
    doc,
    "For most screens, a function is the cleanest UI-to-ViewModel API. Calling it multiple times processes multiple actions, so duplicate network requests should be guarded explicitly rather than assuming that a Channel will remove them."
)
add_code(
    doc,
    """
fun onIntent(intent: MainIntent) {
    when (intent) {
        MainIntent.FetchAnimals -> fetchAnimals()

        is MainIntent.SelectAnimal -> {
            _state.update { it.copy(selectedAnimal = intent.animal) }
        }

        MainIntent.BackToAnimalList -> {
            _state.update { it.copy(selectedAnimal = null) }
        }
    }
}
"""
)
add_body(
    doc,
    "A Channel is useful when actions genuinely need queue semantics, but it should not be added merely because an application is large. Complexity is not architecture. Start with the smallest abstraction that communicates intent clearly."
)

doc.add_heading("5. Fetch data and reduce the result into state", level=1)
add_body(
    doc,
    "The ViewModel launches network work in viewModelScope, emits loading state, calls the repository, and then reduces either the result or error into a new state."
)
add_code(
    doc,
    """
private fun fetchAnimals() {
    if (_state.value.isLoading) return

    viewModelScope.launch {
        _state.update { it.copy(isLoading = true, error = null) }

        try {
            val animals = repository.getAnimals()
            _state.update {
                it.copy(animals = animals, isLoading = false)
            }
        } catch (cancellation: CancellationException) {
            throw cancellation
        } catch (exception: Exception) {
            _state.update {
                it.copy(
                    isLoading = false,
                    error = exception.message ?: "Unable to fetch animals"
                )
            }
        }
    }
}
"""
)
add_body(
    doc,
    "CancellationException deserves special treatment. Coroutine cancellation is control flow, not a user-facing failure. Rethrowing it prevents lifecycle cancellation from incorrectly producing an error message."
)

doc.add_heading("6. Keep the repository focused", level=1)
add_body(
    doc,
    "The repository gives the ViewModel a stable data API and hides where the data comes from. In a small sample it may simply delegate to Retrofit:"
)
add_code(
    doc,
    """
class AnimalRepository @Inject constructor(
    private val api: AnimalApiService
) {
    suspend fun getAnimals(): List<Animal> = api.getAnimals()
}
"""
)
add_body(
    doc,
    "In a larger application, the repository can coordinate a remote API, a Room database, caching rules, and mapping from network models to domain models. The ViewModel does not need to change when those details evolve."
)

doc.add_heading("7. Render state in Compose", level=1)
add_body(
    doc,
    "The Activity obtains the Hilt ViewModel and collects its state in a lifecycle-aware way. Lifecycle-aware collection stops unnecessary observation when the UI is not active."
)
add_code(
    doc,
    """
val state by viewModel.state.collectAsStateWithLifecycle()

AnimalScreen(
    state = state,
    onIntent = viewModel::onIntent
)
"""
)
add_body(
    doc,
    "The screen remains stateless. It receives state and emits intents:"
)
add_code(
    doc,
    """
@Composable
fun AnimalScreen(
    state: MainUiState,
    onIntent: (MainIntent) -> Unit
) {
    when {
        state.isLoading -> AnimalShimmer()

        state.error != null -> ErrorContent(
            message = state.error,
            onRetry = { onIntent(MainIntent.FetchAnimals) }
        )

        state.selectedAnimal != null -> AnimalDetailScreen(
            animal = state.selectedAnimal,
            onBack = { onIntent(MainIntent.BackToAnimalList) }
        )

        else -> AnimalList(
            animals = state.animals,
            onAnimalClick = {
                onIntent(MainIntent.SelectAnimal(it))
            }
        )
    }
}
"""
)

doc.add_heading("8. UI-only state versus feature state", level=1)
add_body(
    doc,
    "Not every value belongs in the ViewModel. A temporary animation progress value or whether a specific image request has finished can remain inside a composable. A selected animal, loading status, fetched list, or error affects the feature’s behavior and should survive recomposition, so it belongs in feature state."
)
add_body(
    doc,
    "A useful test is to ask: if the screen were recreated, would losing this value change what the user believes is happening? If yes, it probably belongs in the ViewModel state."
)

doc.add_heading("StateFlow, SharedFlow, Channel, or function?", level=1)
table = doc.add_table(rows=1, cols=3)
table.autofit = False
widths = [Inches(1.35), Inches(2.35), Inches(2.55)]
headers = ["Tool", "Best direction", "Typical purpose"]
for i, (cell, width, text) in enumerate(zip(table.rows[0].cells, widths, headers)):
    cell.width = width
    set_cell_shading(cell, "DDEBF7")
    set_cell_margins(cell)
    run = cell.paragraphs[0].add_run(text)
    set_run_font(run, bold=True, color=NAVY)
set_repeat_table_header(table.rows[0])

rows = [
    ("Function", "UI → ViewModel", "Send intents with a small, testable API."),
    ("StateFlow", "ViewModel → UI", "Persistent state with a current value."),
    ("SharedFlow", "ViewModel → UI", "Transient effects such as snackbars."),
    ("Channel", "Internal queue", "Sequential one-consumer events when queue semantics are required."),
]
for tool, direction, purpose in rows:
    cells = table.add_row().cells
    for cell, width, text in zip(cells, widths, (tool, direction, purpose)):
        cell.width = width
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(1)

doc.add_heading("Common mistakes", level=1)
for text in (
    "Storing the same list in both a private ViewModel property and UI state.",
    "Letting composables call repositories or Retrofit services directly.",
    "Exposing MutableStateFlow publicly.",
    "Using several independent Boolean flags that can describe impossible combinations.",
    "Treating coroutine cancellation as a user-visible error.",
    "Starting duplicate requests on every repeated tap.",
    "Putting navigation, state reduction, networking, and every composable in one oversized file.",
):
    add_bullet(doc, text)

doc.add_heading("A practical package structure", level=1)
add_code(
    doc,
    """
com.example.animals
├── data
│   ├── api
│   │   └── AnimalApiService.kt
│   ├── model
│   │   └── Animal.kt
│   └── repository
│       └── AnimalRepository.kt
├── di
│   └── NetworkModule.kt
├── feature
│   └── animals
│       ├── MainIntent.kt
│       ├── MainUiState.kt
│       ├── MainViewModel.kt
│       ├── AnimalScreen.kt
│       └── components
└── MviApplication.kt
"""
)
add_body(
    doc,
    "Package by feature once the application grows. The goal is not to create the maximum number of layers; it is to keep ownership clear and changes local."
)

doc.add_heading("Testing becomes straightforward", level=1)
add_body(
    doc,
    "Because the ViewModel accepts intents and exposes state, tests can read like behavior specifications: send FetchAnimals, verify loading, return repository data, and verify the final list. Send SelectAnimal and verify the selected animal. Send BackToAnimalList and verify selection is cleared."
)
add_body(
    doc,
    "Introducing a repository interface or a fake repository makes these tests independent from Retrofit and the network."
)

doc.add_heading("When MVI is a good fit", level=1)
add_body(
    doc,
    "MVI is especially useful for screens with asynchronous data, multiple user actions, loading and retry states, selection, filtering, pagination, or navigation-like transitions. For a tiny static screen, a complete intent/state/reducer setup may be unnecessary."
)
add_callout(
    doc,
    "The real goal",
    "Do not adopt MVI to collect architecture terminology. Adopt it when one-way data flow makes screen behavior easier to understand, test, and change."
)

doc.add_heading("Final takeaway", level=1)
add_body(
    doc,
    "A solid Android MVI implementation has a small number of non-negotiable properties: the UI sends actions rather than mutating data, the ViewModel is the state owner, state is immutable and observable, data access is delegated to a repository, and the screen renders from one source of truth."
)
add_body(
    doc,
    "Once those boundaries are in place, Retrofit, Hilt, coroutines, Compose, Glide, and shimmer placeholders become implementation details around a predictable core. That predictability—not the number of classes—is what makes the architecture valuable."
)

doc.add_heading("Medium publishing notes", level=1)
add_body(
    doc,
    "Suggested subtitle: A real-world walkthrough of one-way data flow, immutable UI state, and predictable Compose screens."
)
add_body(
    doc,
    "Suggested tags: Android, Kotlin, Jetpack Compose, Software Architecture, Mobile Development."
)

# Avoid widows where possible.
for paragraph in doc.paragraphs:
    set_paragraph_keep(paragraph, keep_lines=True)

doc.core_properties.title = "A Practical Guide to MVI Architecture in Android"
doc.core_properties.subject = "Medium-ready Android architecture article"
doc.core_properties.author = "Neha"
doc.core_properties.keywords = "Android, Kotlin, MVI, Jetpack Compose, StateFlow, Hilt"

doc.save(OUTPUT)
print(OUTPUT)
